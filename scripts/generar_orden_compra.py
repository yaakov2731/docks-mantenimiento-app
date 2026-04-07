#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from datetime import date

# ── DATA ─────────────────────────────────────────────────────────────────────
pedidos = [
    {"id": "PED782923", "fecha": "3/4/2026", "local": "Parrilla",       "responsable": "Umo",    "semana": "Semana 14", "prioridad": "Urgente"},
    {"id": "PED218650", "fecha": "3/4/2026", "local": "Cafetería",      "responsable": "Angie",  "semana": "Semana 14", "prioridad": "Urgente"},
    {"id": "PED386295", "fecha": "4/4/2026", "local": "Parrilla",       "responsable": "Umo",    "semana": "Semana 14", "prioridad": "Urgente"},
    {"id": "PED404445", "fecha": "5/4/2026", "local": "Parrilla",       "responsable": "Ayelén", "semana": "Semana 15", "prioridad": "Normal"},
    {"id": "PED519309", "fecha": "5/4/2026", "local": "Parrilla",       "responsable": "Ayelén", "semana": "Semana 15", "prioridad": "Normal"},
    {"id": "PED369749", "fecha": "5/4/2026", "local": "Hamburguesería", "responsable": "Maru",   "semana": "Semana 15", "prioridad": "Urgente"},
    {"id": "PED975820", "fecha": "5/4/2026", "local": "Heladería",      "responsable": "Kiki",   "semana": "Semana 14", "prioridad": "Urgente"},
    {"id": "PED444777", "fecha": "5/4/2026", "local": "Cafetería",      "responsable": "Angie",  "semana": "Semana 15", "prioridad": "Urgente"},
]

# Consolidated items by category with source locals
categorias = {
    "VERDURAS Y FRUTAS FRESCAS": [
        ("Rúcula",              "1 caja + 17 paquetes", "Parrilla, Cafetería, Hamburguesería"),
        ("Morrón rojo",         "2 kg",                 "Parrilla"),
        ("Tomate",              "1 caja",               "Parrilla"),
        ("Lechuga",             "1 caja + 2 kg",        "Parrilla, Hamburguesería"),
        ("Cebolla morada",      "3 kg",                 "Parrilla"),
        ("Cebolla",             "3 kg",                 "Hamburguesería"),
        ("Ajo",                 "1 unidad",             "Parrilla"),
        ("Perejil",             "2 paquetes",           "Parrilla"),
        ("Cherrys",             "1 kg",                 "Cafetería"),
        ("Naranja",             "1 caja",               "Cafetería"),
        ("Limón medio",         "1 caja",               "Cafetería"),
        ("Palta",               "3 unidades",           "Cafetería"),
        ("Banana",              "3 kg",                 "Cafetería, Heladería"),
        ("Frutilla fresca",     "500 g",                "Cafetería"),
        ("Manzana roja",        "1 kg",                 "Cafetería"),
        ("Kiwi",                "1 kg",                 "Cafetería"),
        ("Menta",               "1 bolsa + 2 paquetes", "Heladería, Cafetería"),
    ],
    "FRUTAS CONGELADAS": [
        ("Frutilla congelada",  "2 bolsas",             "Heladería"),
        ("Mango congelado",     "2 bolsas",             "Heladería"),
    ],
    "LÁCTEOS, HUEVOS Y QUESOS": [
        ("Leche",               "11 l + 2 cajas",       "Parrilla, Heladería, Cafetería"),
        ("Crema de leche",      "1,5 l",                "Parrilla"),
        ("Crema en pomo",       "1 unidad",             "Cafetería"),
        ("Queso crema",         "3 l",                  "Cafetería"),
        ("Queso Sardo",         "1 unidad",             "Parrilla"),
        ("Queso Tybo",          "1 unidad",             "Hamburguesería"),
        ("Huevos",              "150 unidades",         "Hamburguesería, Cafetería"),
    ],
    "CARNES, FIAMBRES Y EMBUTIDOS": [
        ("Milanesas",           "5 kg",                 "Hamburguesería"),
        ("Pechuga de pollo",    "5 kg",                 "Hamburguesería"),
        ("Suprema de pollo",    "5 kg",                 "Parrilla"),
        ("Filet",               "1 paquete",            "Parrilla"),
        ("Rabas",               "2 paquetes",           "Parrilla"),
        ("Panceta",             "1 unidad",             "Hamburguesería"),
        ("Jamón",               "1 unidad",             "Cafetería"),
        ("Salchichas",          "2 paquetes",           "Heladería"),
    ],
    "ACEITES, SALSAS Y CONDIMENTOS": [
        ("Aceite",              "65 l",                 "Parrilla, Hamburguesería"),
        ("Aceite de oliva",     "1 unidad",             "Parrilla"),
        ("Salsa caesar",        "3 unidades",           "Parrilla"),
        ("Salsa de dulce de leche", "4 unidades",       "Heladería, Cafetería"),
        ("Salsa caramelo",      "2 unidades",           "Cafetería"),
        ("Triturado",           "1 paquete",            "Parrilla"),
    ],
    "PANIFICADOS Y HARINAS": [
        ("Harina",              "4 kg",                 "Parrilla"),
        ("Pan rallado / Rayado","2 paquetes",           "Parrilla, Hamburguesería"),
        ("Tapas de empanadas",  "1 paquete",            "Parrilla"),
        ("Fideos (personal)",   "5 paquetes",           "Parrilla"),
        ("Pan de lomo",         "3 kg",                 "Hamburguesería"),
        ("Pan de hamburguesas", "40 unidades",          "Hamburguesería"),
        ("Pan de pancho",       "18 unidades + 2 paquetes", "Hamburguesería, Heladería"),
        ("Pan de miga",         "2 unidades",           "Heladería, Cafetería"),
        ("Pan árabe",           "10 unidades",          "Cafetería"),
    ],
    "AZÚCAR, DULCES Y REPOSTERÍA": [
        ("Azúcar",              "20 kg",                "Heladería, Cafetería"),
        ("Dulce de leche repostero (Vacalín)", "4 kg",  "Cafetería"),
        ("Chips de chocolate negro y blanco", "100 g",  "Cafetería"),
        ("Rockets",             "100 g",                "Cafetería"),
        ("Chocolate amargo cobertura", "1 unidad",      "Cafetería"),
        ("Mermelada de frutilla","1 unidad",            "Cafetería"),
        ("Gelatina sin sabor",  "2 unidades",           "Cafetería"),
        ("Aros de cebolla (prep.)","2 paquetes",        "Hamburguesería"),
    ],
    "HELADOS": [
        ("Helado Chocolate con almendras", "1 unidad",  "Heladería"),
        ("Helado Flan con dulce de leche", "1 unidad",  "Heladería"),
        ("Helado Banana Split", "1 unidad",             "Heladería"),
        ("Helado Tramontana",   "1 unidad",             "Heladería"),
        ("Helado Mousse de chocolate", "2 unidades",    "Heladería"),
    ],
    "BEBIDAS": [
        ("Agua sin gas",        "4 paquetes",           "Heladería"),
        ("Agua con gas",        "4 paquetes",           "Heladería"),
        ("Coca-Cola",           "1 caja",               "Heladería"),
        ("Coca-Cola Zero",      "1 caja",               "Heladería"),
        ("Sprite",              "1 caja",               "Heladería"),
        ("Sprite Zero",         "1 caja",               "Heladería"),
        ("Fanta",               "1 caja",               "Heladería"),
        ("Fanta Zero",          "1 caja",               "Heladería"),
        ("Aquarius Pera",       "2 paquetes",           "Heladería"),
        ("Aquarius Pomelo",     "2 paquetes",           "Heladería"),
    ],
    "ENLATADOS Y CONSERVAS": [
        ("Atún en lata",        "3 latas",              "Cafetería"),
    ],
    "INSUMOS, LIMPIEZA Y DESCARTABLES": [
        ("Aceite — aceites en envase (no comestible)", "—", "—"),  # placeholder remove
        ("Rejillas",            "6 unidades",           "Parrilla, Heladería"),
        ("Fibra parrillera",    "1 paquete",            "Parrilla"),
        ("Secador de verdura",  "1 unidad",             "Parrilla"),
        ("Bobina",              "1 unidad",             "Heladería"),
        ("Trapo de piso",       "2 unidades",           "Heladería"),
        ("Guantes de nitrilo",  "1 caja",               "Heladería"),
        ("Potes de ¼",          "2 bolsas",             "Heladería"),
        ("Vasos de 500 ml",     "2 bolsas",             "Heladería"),
        ("Cucuruchos",          "1 caja",               "Heladería"),
    ],
}

# Remove placeholder
categorias["INSUMOS, LIMPIEZA Y DESCARTABLES"] = [
    r for r in categorias["INSUMOS, LIMPIEZA Y DESCARTABLES"]
    if r[0] != "Aceite — aceites en envase (no comestible)"
]

# ── HELPERS ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), kwargs.get('val', 'single'))
        tag.set(qn('w:sz'), kwargs.get('sz', '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', '000000'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def bold_run(para, text, size=11, color=None):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return run

def normal_run(para, text, size=10):
    run = para.add_run(text)
    run.font.size = Pt(size)
    return run

# ── BUILD DOCUMENT ────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── HEADER ────────────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
bold_run(title_para, "ORDEN DE COMPRA CONSOLIDADA", size=18, color="1B4F72")

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
bold_run(sub_para, "Docks Gastronomía  ·  Semanas 14 & 15 — Abril 2026", size=12, color="555555")

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
normal_run(date_para, f"Generado: {date.today().strftime('%d/%m/%Y')}    |    Prioridad: URGENTE + NORMAL")

doc.add_paragraph()

# ── RESUMEN DE PEDIDOS ────────────────────────────────────────────────────────
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.LEFT
bold_run(h, "RESUMEN DE PEDIDOS INCLUIDOS", size=12, color="1B4F72")

tbl = doc.add_table(rows=1, cols=5)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr_cells = tbl.rows[0].cells
headers = ["N° Pedido", "Fecha", "Local", "Responsable", "Prioridad"]
widths = [Cm(3), Cm(2.8), Cm(4), Cm(3.5), Cm(2.8)]
for i, (cell, h_text, w) in enumerate(zip(hdr_cells, headers, widths)):
    cell.width = w
    set_cell_bg(cell, "1B4F72")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h_text)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# Data rows
for idx, p in enumerate(pedidos):
    row = tbl.add_row().cells
    row_data = [p["id"], p["fecha"], p["local"], p["responsable"], p["prioridad"]]
    bg = "EBF5FB" if idx % 2 == 0 else "FFFFFF"
    for cell, val in zip(row, row_data):
        set_cell_bg(cell, bg)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(val)
        run.font.size = Pt(9)
        if val == "Urgente":
            run.bold = True
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

doc.add_paragraph()

# ── ORDEN CONSOLIDADA POR CATEGORÍA ──────────────────────────────────────────
h2 = doc.add_paragraph()
bold_run(h2, "DETALLE CONSOLIDADO POR CATEGORÍA", size=12, color="1B4F72")

CAT_COLORS = {
    "VERDURAS Y FRUTAS FRESCAS":            "27AE60",
    "FRUTAS CONGELADAS":                    "1ABC9C",
    "LÁCTEOS, HUEVOS Y QUESOS":             "F39C12",
    "CARNES, FIAMBRES Y EMBUTIDOS":         "C0392B",
    "ACEITES, SALSAS Y CONDIMENTOS":        "8E44AD",
    "PANIFICADOS Y HARINAS":                "D4AC0D",
    "AZÚCAR, DULCES Y REPOSTERÍA":          "E91E8C",
    "HELADOS":                              "2980B9",
    "BEBIDAS":                              "16A085",
    "ENLATADOS Y CONSERVAS":               "7F8C8D",
    "INSUMOS, LIMPIEZA Y DESCARTABLES":     "555555",
}

for cat_name, items in categorias.items():
    cat_color = CAT_COLORS.get(cat_name, "1B4F72")

    # Category header table (single row, full width)
    cat_tbl = doc.add_table(rows=1, cols=1)
    cat_tbl.style = 'Table Grid'
    cat_cell = cat_tbl.rows[0].cells[0]
    cat_cell.width = Cm(16)
    set_cell_bg(cat_cell, cat_color)
    cp = cat_cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cr = cp.add_run(f"  {cat_name}")
    cr.bold = True
    cr.font.size = Pt(10)
    cr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Items table
    item_tbl = doc.add_table(rows=1, cols=3)
    item_tbl.style = 'Table Grid'

    # Column headers
    col_hdr = item_tbl.rows[0].cells
    for cell, txt, w in zip(col_hdr,
                             ["Producto", "Cantidad Total", "Locales Solicitantes"],
                             [Cm(7), Cm(3.5), Cm(5.5)]):
        cell.width = w
        set_cell_bg(cell, "D5D8DC")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(9)

    for i, (producto, cantidad, locales) in enumerate(items):
        row = item_tbl.add_row().cells
        bg = "F9F9F9" if i % 2 == 0 else "FFFFFF"
        for cell, val, align in zip(row,
                                    [producto, cantidad, locales],
                                    [WD_ALIGN_PARAGRAPH.LEFT,
                                     WD_ALIGN_PARAGRAPH.CENTER,
                                     WD_ALIGN_PARAGRAPH.LEFT]):
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = align
            r = p.add_run(val)
            r.font.size = Pt(9)

    doc.add_paragraph()

# ── FOOTER NOTE ──────────────────────────────────────────────────────────────
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run("Documento generado automáticamente — Docks Gastronomía · Sistema de Gestión de Pedidos")
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
r.italic = True

# ── SAVE ─────────────────────────────────────────────────────────────────────
output = "/home/user/docks-mantenimiento-app/Orden_Compra_Consolidada_Semana14-15_Abril2026.docx"
doc.save(output)
print(f"Documento guardado: {output}")
