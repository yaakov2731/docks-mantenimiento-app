from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "Manual_Empleados_Gastronomia_Asistencia_Docks.docx"

INK = RGBColor(31, 35, 40)
CHARCOAL = RGBColor(39, 43, 48)
GOLD = RGBColor(174, 129, 46)
MUTED = RGBColor(101, 109, 118)
BLUE_GRAY = RGBColor(232, 238, 245)
LIGHT_GOLD = "FFF7E6"
LIGHT_GRAY = "F6F7F9"
RULE = "D9DEE7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=RULE, size="6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            set_cell_margins(cell)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_run(paragraph, text, bold=False, color=INK, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def add_para(doc, text="", style=None, before=0, after=6, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    if align is not None:
        p.alignment = align
    if text:
        add_run(p, text)
    return p


def add_heading(doc, text, level=1):
    style = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_with_next = True
    add_run(p, text, bold=True, color=CHARCOAL if level == 1 else GOLD, size=16 if level == 1 else 13)
    return p


def add_check(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        add_run(p, item)


def add_note_box(doc, title, body, fill=LIGHT_GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, color="E4C77B", size="8")
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title, bold=True, color=CHARCOAL, size=11.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    add_run(p2, body, color=INK, size=10.5)
    add_para(doc, after=8)


def add_flow_table(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_widths(table, [1.55, 1.65, 1.65, 1.65])
    labels = [
        ("1", "Enviar mensaje", "Abrir el bot de WhatsApp."),
        ("2", "Elegir asistencia", "Entrar al menu de fichada."),
        ("3", "Registrar accion", "Entrada, almuerzo o salida."),
        ("4", "Revisar respuesta", "Confirmar que el bot diga registrado."),
    ]
    for idx, (num, title, body) in enumerate(labels):
        cell = table.cell(0, idx)
        set_cell_shading(cell, "FFFFFF" if idx % 2 == 0 else LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, num, bold=True, color=GOLD, size=20)
        t = cell.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        t.paragraph_format.space_after = Pt(2)
        add_run(t, title, bold=True, color=CHARCOAL, size=10.5)
        b = cell.add_paragraph()
        b.alignment = WD_ALIGN_PARAGRAPH.CENTER
        b.paragraph_format.space_after = Pt(0)
        add_run(b, body, color=MUTED, size=9.5)
    add_para(doc, after=8)


def add_quick_table(doc):
    rows = [
        ("Entrada", "Cuando llegas al local y estas listo para empezar.", "Marca inicio del dia."),
        ("Inicio de almuerzo", "Cuando paras para almorzar.", "El tiempo deja de contar como trabajo."),
        ("Fin de almuerzo", "Cuando volves a trabajar.", "El sistema retoma el conteo."),
        ("Salida", "Cuando termina tu turno.", "Cierra el registro del dia."),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_widths(table, [1.65, 3.05, 1.8])
    headers = ["Accion", "Cuando usarla", "Resultado"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        add_run(p, h, bold=True, color=CHARCOAL, size=10.5)
    set_repeat_table_header(table.rows[0])
    for action, when, result in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((action, when, result)):
            set_cell_shading(cells[idx], "FFFFFF")
            p = cells[idx].paragraphs[0]
            add_run(p, text, bold=idx == 0, color=INK, size=10)
    add_para(doc, after=8)


def add_problem_table(doc):
    rows = [
        ("El bot no responde", "Esperar un minuto y reenviar 'menu'. Si sigue igual, avisar al encargado."),
        ("Me equivoque de opcion", "Enviar '0' para volver o 'menu' para empezar de nuevo."),
        ("No aparece mi nombre", "Avisar al encargado. El numero debe estar dado de alta."),
        ("Fiche mal", "No intentar arreglarlo con mensajes repetidos. Avisar al encargado."),
        ("Cambie de local", "Avisar antes de fichar para que el registro entre al local correcto."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_widths(table, [2.0, 4.5])
    headers = ["Situacion", "Que hacer"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "E8EEF5")
        add_run(cell.paragraphs[0], h, bold=True, color=CHARCOAL, size=10.5)
    set_repeat_table_header(table.rows[0])
    for situation, action in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], "FFFFFF")
        set_cell_shading(cells[1], "FFFFFF")
        add_run(cells[0].paragraphs[0], situation, bold=True, color=INK, size=10)
        add_run(cells[1].paragraphs[0], action, color=INK, size=10)
    add_para(doc, after=8)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for style_name in ("Heading 1", "Heading 2"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.bold = True
        style.paragraph_format.keep_with_next = True

    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = CHARCOAL
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)

    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = GOLD
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header, "Docks del Puerto | Guia de asistencia gastronomia", color=MUTED, size=8.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "Uso interno - Empleados de gastronomia", color=MUTED, size=8.5)

    p = add_para(doc, "GUIA RAPIDA PARA EMPLEADOS", before=8, after=3)
    add_run(p, "", color=GOLD)
    p.runs[0].font.color.rgb = GOLD
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(10)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    add_run(title, "Como fichar asistencia por WhatsApp", bold=True, color=CHARCOAL, size=25)

    subtitle = add_para(
        doc,
        "Docks del Puerto - Polo Gastronomico | Entrada, almuerzo, salida y registro en planilla",
        after=12,
    )
    subtitle.runs[0].font.color.rgb = MUTED
    subtitle.runs[0].font.size = Pt(11.5)

    meta = add_para(doc, f"Version: {date.today().strftime('%d/%m/%Y')}  |  Documento interno", after=18)
    meta.runs[0].font.color.rgb = GOLD
    meta.runs[0].font.bold = True
    meta.runs[0].font.size = Pt(10)

    add_note_box(
        doc,
        "Objetivo",
        "Que cada empleado registre su asistencia desde WhatsApp de forma simple y correcta. "
        "El registro alimenta automaticamente la planilla de asistencia y las hojas de sueldos del local.",
    )

    add_heading(doc, "El circuito completo", 1)
    add_flow_table(doc)

    add_heading(doc, "Que tenes que hacer cada dia", 1)
    add_numbered(
        doc,
        [
            "Al llegar al local, escribi al bot y entra en Registrar asistencia.",
            "Elegi Registrar entrada y espera la confirmacion del bot.",
            "Si paras para almorzar, marca Inicio de almuerzo antes de cortar.",
            "Cuando volves del almuerzo, marca Fin de almuerzo.",
            "Al terminar tu turno, marca Registrar salida.",
        ],
    )

    add_note_box(
        doc,
        "Regla principal",
        "No fiches dos veces la misma accion. Si te equivocaste, avisale al encargado para corregirlo desde administracion.",
        fill="FFF2CC",
    )

    add_heading(doc, "Acciones del bot", 1)
    add_quick_table(doc)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Como usar el bot sin confundirte", 1)
    add_check(doc, "Usa numeros cuando el bot muestra opciones.")
    add_check(doc, "Si queres volver, manda 0.")
    add_check(doc, "Si te perdiste, manda menu.")
    add_check(doc, "Lee la respuesta final del bot. Tiene que decir que quedo registrado.")
    add_check(doc, "Si trabajas en otro local ese dia, avisa antes de fichar.")

    add_heading(doc, "Que pasa despues de fichar", 1)
    add_para(
        doc,
        "Cuando fichas, el sistema guarda el movimiento en la app y sincroniza la planilla de Google Sheets.",
    )
    add_check(doc, "La pestaña Asistencia_App recibe el registro real.")
    add_check(doc, "La pestaña de sueldos del local marca el dia trabajado si el empleado y el local coinciden.")
    add_check(doc, "Planificacion queda solo para organizar la semana, no para liquidar.")

    add_note_box(
        doc,
        "Importante para el pago",
        "Tu nombre debe estar bien cargado. Si el nombre esta distinto en la planilla, el dia puede no marcarse automaticamente.",
        fill="EAF4EA",
    )

    add_heading(doc, "Errores comunes", 1)
    add_problem_table(doc)

    add_heading(doc, "Checklist antes de irte", 1)
    add_check(doc, "Registre mi salida.")
    add_check(doc, "El bot confirmo que la salida quedo registrada.")
    add_check(doc, "Avise cualquier error de fichada al encargado.")
    add_check(doc, "No mande mensajes repetidos para corregir una fichada.")

    add_note_box(
        doc,
        "Contacto interno",
        "Ante dudas o errores de asistencia, hablar con el encargado del local o administracion antes del cierre de la jornada.",
        fill=LIGHT_GRAY,
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_doc()
    print(DOCX_PATH)
